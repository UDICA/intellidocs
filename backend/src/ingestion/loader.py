"""Multi-format document loader for the ingestion pipeline.

Supports TXT, Markdown, CSV, PDF, and DOCX file formats. Each file is loaded
into one or more Document objects carrying the raw text content and metadata
(source path, format, page number where applicable).
"""

from __future__ import annotations

import csv
import io
import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".pdf", ".docx"}


@dataclass
class Document:
    """A loaded document with content and metadata."""

    content: str
    metadata: dict[str, str | int] = field(default_factory=dict)


class DocumentLoader:
    """Load documents from various file formats.

    Provides a unified interface for reading files into Document objects.
    PDF and DOCX support requires ``pypdf`` and ``python-docx`` respectively;
    these are imported lazily so the loader can be used for plain-text
    formats without those heavy dependencies installed.
    """

    def load(self, path: Path) -> list[Document]:
        """Load a single file and return a list of Documents.

        Args:
            path: Path to the file to load.

        Returns:
            A list of Document objects. Most formats produce a single
            document; PDFs produce one per page.

        Raises:
            FileNotFoundError: If *path* does not exist.
            ValueError: If the file extension is not in SUPPORTED_EXTENSIONS.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}")

        loader = {
            ".txt": self._load_text,
            ".md": self._load_text,
            ".markdown": self._load_text,
            ".csv": self._load_csv,
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
        }[ext]

        docs = loader(path)
        for doc in docs:
            doc.metadata.setdefault("source", str(path))
            doc.metadata.setdefault("format", ext.lstrip("."))
        return docs

    def load_directory(self, directory: Path) -> list[Document]:
        """Load all supported files from a directory (non-recursive).

        Files are processed in sorted order for deterministic output.
        Failures on individual files are logged and skipped so that one
        corrupt file does not abort the entire batch.

        Args:
            directory: Path to the directory to scan.

        Returns:
            A flat list of Document objects from all successfully loaded files.
        """
        directory = Path(directory)
        docs: list[Document] = []
        for path in sorted(directory.iterdir()):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    docs.extend(self.load(path))
                except Exception:
                    logger.warning("Failed to load %s", path, exc_info=True)
        return docs

    # ------------------------------------------------------------------
    # Private format-specific loaders
    # ------------------------------------------------------------------

    def _load_text(self, path: Path) -> list[Document]:
        """Load a plain-text or Markdown file."""
        content = path.read_text(encoding="utf-8")
        return [Document(content=content)]

    def _load_csv(self, path: Path) -> list[Document]:
        """Load a CSV file, combining all rows into a single document.

        Each row is rendered as ``key: value | key: value | ...`` so that
        the downstream chunker receives human-readable text rather than
        raw delimited data.
        """
        content = path.read_text(encoding="utf-8")
        reader = csv.DictReader(io.StringIO(content))
        rows = [
            " | ".join(f"{k}: {v}" for k, v in row.items())
            for row in reader
        ]
        combined = "\n".join(rows)
        return [Document(content=combined)]

    def _load_pdf(self, path: Path) -> list[Document]:
        """Load a PDF file, producing one Document per page."""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        docs: list[Document] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    content=text,
                    metadata={"page": i + 1},
                ))
        return docs

    def _load_docx(self, path: Path) -> list[Document]:
        """Load a DOCX file, joining non-empty paragraphs."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return [Document(content=content)]
